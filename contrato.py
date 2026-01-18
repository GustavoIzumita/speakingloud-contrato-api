from flask import Flask, request, send_file, jsonify, make_response
from flask_cors import CORS
from docx import Document
from docx.shared import Pt
from datetime import date
import uuid
import os
import traceback

app = Flask(__name__)
CORS(app)

def pegar(dados, chave):
    return str(dados.get(chave, ""))

@app.route("/gerar-contrato", methods=["POST", "OPTIONS"])
def gerar_contrato():

    if request.method == "OPTIONS":
        response = make_response()
        response.headers["Access-Control-Allow-Origin"] = "*"
        response.headers["Access-Control-Allow-Headers"] = "Content-Type"
        response.headers["Access-Control-Allow-Methods"] = "POST, OPTIONS"
        return response, 200

    try:
        dados = request.json or {}

        # DATA AUTOMÁTICA
        hoje = date.today()
        meses = {
            1: "janeiro", 2: "fevereiro", 3: "março",
            4: "abril", 5: "maio", 6: "junho",
            7: "julho", 8: "agosto", 9: "setembro",
            10: "outubro", 11: "novembro", 12: "dezembro"
        }
        data_formatada = f"{hoje.day} de {meses[hoje.month]} de {hoje.year}"

        doc = Document("contrato_modelo.docx")

        campos = {
            "{{NOME}}": pegar(dados, "nome"),
            "{{CPF}}": pegar(dados, "cpf"),
            "{{ENDERECO}}": pegar(dados, "endereco"),
            "{{BAIRRO}}": pegar(dados, "bairro"),
            "{{CEP}}": pegar(dados, "cep"),
            "{{CIDADE}}": pegar(dados, "cidade"),
            "{{UF}}": pegar(dados, "uf"),
            "{{TELEFONE}}": pegar(dados, "telefone"),
            "{{EMAIL}}": pegar(dados, "email"),
            "{{ALUNO}}": pegar(dados, "aluno"),
            "{{CURSO}}": pegar(dados, "curso"),
            "{{MODALIDADE}}": pegar(dados, "modalidade"),
            "{{VALOR}}": pegar(dados, "valor"),
            "{{DATA}}": data_formatada
        }

        # PARÁGRAFOS
        for paragrafo in doc.paragraphs:
            for run in paragrafo.runs:
                for campo, valor in campos.items():
                    if campo in run.text:
                        run.text = run.text.replace(campo, valor)
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

        # TABELAS
        for tabela in doc.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        for run in paragrafo.runs:
                            for campo, valor in campos.items():
                                if campo in run.text:
                                    run.text = run.text.replace(campo, valor)
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(12)
                            import unicodedata  # coloque no topo do arquivo, junto com os imports
                            nome = pegar(dados, "nome")
                            
                            # remove acentos (João -> Joao
                            nome_sem_acento = unicodedata.normalize("NFKD", nome).encode("ASCII", "ignore").decode("ASCII")
                            
                            # remove espaços extra
                            nome_base = nome_sem_acento.strip().replace(" ", "_")
                            
                            arquivo_docx = f"Contrato-{nome_base}.docx"
                            
                            doc.save(arquivo_docx)

        response = send_file(
            arquivo_docx,
            as_attachment=True,
            download_name=arquivo_docx
        )
        response.headers["Access-Control-Allow-Origin"] = "*"
        return response

    except Exception as e:
        traceback.print_exc()
        return jsonify({"erro": str(e)}), 500

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
